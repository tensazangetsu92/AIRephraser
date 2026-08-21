import io
import json
import logging
import re
from copy import deepcopy
from time import perf_counter

from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app.database import User, get_db
from app.models import HumanizeRequest
from app.subscription import (
    check_usage_limit,
    check_study_work_request_limit,
    complete_study_work_trial,
    count_words,
    get_user_subscription,
    increment_usage,
    release_study_work_trial,
    reserve_study_work_trial,
)
from auth import get_current_user
from llm import LLMProcessingError, LLMTimeoutError, humanize_docx_paragraphs, humanize_pipeline

router = APIRouter(tags=["study-work"])
logger = logging.getLogger("uvicorn.error")

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_DOCX_SIZE_BYTES = 10 * 1024 * 1024


def raise_llm_error(error: Exception) -> None:
    if isinstance(error, LLMTimeoutError):
        raise HTTPException(status_code=504, detail="Нейросеть не ответила за 60 секунд. Попробуйте ещё раз.")
    raise HTTPException(status_code=502, detail="Нейросеть временно недоступна. Попробуйте ещё раз немного позже.")


def ensure_text_not_empty(text: str) -> None:
    if not text or not text.strip():
        raise HTTPException(status_code=400, detail="В документе не найден текст для обработки.")


def reserve_trial_or_require_paid_plan(db: Session, user_id: int) -> bool:
    """Free users can use Study Work only through their one registration trial."""
    uses_trial = reserve_study_work_trial(db, user_id)
    if uses_trial:
        return True
    if get_user_subscription(db, user_id).plan_type == "free":
        raise HTTPException(
            status_code=403,
            detail="Бесплатное использование «Учебной работы» уже использовано. Выберите платный тариф, чтобы продолжить.",
        )
    return False


def is_editable_docx_paragraph(paragraph) -> bool:
    if not paragraph.text.strip():
        return False
    style_name = (getattr(paragraph.style, "name", "") or "").lower()
    protected_style_parts = ("heading", "title", "заголов", "подзаголов", "название", "toc")
    return not any(part in style_name for part in protected_style_parts)


def is_bibliography_heading(text: str) -> bool:
    normalized = re.sub(r"\s+", " ", text).strip().casefold().rstrip(".:")
    headings = (
        "список использованных источников",
        "список литературы",
        "библиографический список",
        "references",
        "bibliography",
    )
    return normalized in headings


def is_numbered_docx_paragraph(paragraph) -> bool:
    if re.match(r"^\s*(?:\d+|[IVXLCDM]+)[.)]\s+", paragraph.text, flags=re.IGNORECASE):
        return True
    paragraph_properties = getattr(paragraph._p, "pPr", None)
    return bool(paragraph_properties is not None and paragraph_properties.numPr is not None)


def has_bibliography_markers(text: str) -> bool:
    normalized = text.casefold()
    return (
        "http://" in normalized
        or "https://" in normalized
        or "doi:" in normalized
        or "doi.org/" in normalized
        or "дата обращения" in normalized
        or "retrieved" in normalized
        or "accessed" in normalized
    )


def find_bibliography_start(paragraphs: list) -> int | None:
    for index, paragraph in enumerate(paragraphs):
        if is_bibliography_heading(paragraph.text):
            return index
    for index in range(len(paragraphs) - 1):
        current, following = paragraphs[index], paragraphs[index + 1]
        if not (is_numbered_docx_paragraph(current) and is_numbered_docx_paragraph(following)):
            continue
        if has_bibliography_markers(current.text) or has_bibliography_markers(following.text):
            return index
    return None


def get_editable_docx_paragraphs(document) -> list:
    paragraphs = document.paragraphs
    bibliography_start = find_bibliography_start(paragraphs)
    return [
        paragraph
        for index, paragraph in enumerate(paragraphs)
        if (bibliography_start is None or index < bibliography_start)
        and is_editable_docx_paragraph(paragraph)
    ]


def replace_docx_paragraph_text(paragraph, replacement: str) -> None:
    first_run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        first_run_properties = deepcopy(paragraph.runs[0]._r.rPr)
    paragraph.clear()
    new_run = paragraph.add_run(replacement)
    if first_run_properties is not None:
        new_run._r.insert(0, first_run_properties)


@router.post("/study-work/process")
async def process_study_work(
        req: HumanizeRequest,
        current_user: User = Depends(get_current_user),
        db: Session = Depends(get_db),
):
    ensure_text_not_empty(req.text)
    word_count = count_words(req.text)
    study_work_limit = check_study_work_request_limit(db, current_user.id, req.text)
    uses_trial = reserve_trial_or_require_paid_plan(db, current_user.id)
    if not uses_trial:
        check_usage_limit(db, current_user.id, req.text, max_words_override=study_work_limit)

    started_at = perf_counter()
    try:
        result = await humanize_pipeline(
            req.text, "academic", "professional", req.target_language,
            req.academic_work_type, req.preserve_options,
        )
    except (LLMTimeoutError, LLMProcessingError) as error:
        if uses_trial:
            release_study_work_trial(db, current_user.id)
        raise_llm_error(error)
    except Exception:
        if uses_trial:
            release_study_work_trial(db, current_user.id)
        raise

    if uses_trial:
        complete_study_work_trial(db, current_user.id)
    else:
        increment_usage(db, current_user.id, word_count)
    logger.info("study_work: %s words processed in %.2f seconds", word_count, perf_counter() - started_at)
    return {"success": True, "result": result, "used_free_trial": uses_trial}


@router.post("/study-work/docx/process")
async def process_study_work_docx(
        file: UploadFile = File(...),
        work_type: str = Form("other"),
        preserve_options: str = Form("[]"),
        current_user: User = Depends(get_current_user),
        db: Session = Depends(get_db),
):
    filename = file.filename or ""
    if not filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Загрузите документ Word в формате .docx.")
    document_bytes = await file.read()
    if not document_bytes:
        raise HTTPException(status_code=400, detail="Файл Word пустой.")
    if len(document_bytes) > MAX_DOCX_SIZE_BYTES:
        raise HTTPException(status_code=413, detail="Файл Word больше 10 МБ.")
    try:
        selected_options = json.loads(preserve_options)
        if not isinstance(selected_options, list):
            selected_options = []
    except json.JSONDecodeError:
        selected_options = []
    try:
        document = Document(io.BytesIO(document_bytes))
    except Exception as error:
        logger.warning("Unable to open uploaded DOCX: %s", error)
        raise HTTPException(status_code=400, detail="Не удалось открыть файл Word.")

    editable_paragraphs = get_editable_docx_paragraphs(document)
    source_text = "\n\n".join(paragraph.text for paragraph in editable_paragraphs)
    ensure_text_not_empty(source_text)
    word_count = count_words(source_text)
    study_work_limit = check_study_work_request_limit(db, current_user.id, source_text)
    uses_trial = reserve_trial_or_require_paid_plan(db, current_user.id)
    if not uses_trial:
        check_usage_limit(db, current_user.id, source_text, max_words_override=study_work_limit)
    started_at = perf_counter()
    items = [{"id": f"p{index}", "text": paragraph.text} for index, paragraph in enumerate(editable_paragraphs)]
    try:
        replacements = await humanize_docx_paragraphs(items, selected_options, work_type)
    except (LLMTimeoutError, LLMProcessingError) as error:
        if uses_trial:
            release_study_work_trial(db, current_user.id)
        raise_llm_error(error)
    except Exception:
        if uses_trial:
            release_study_work_trial(db, current_user.id)
        raise

    for item, paragraph in zip(items, editable_paragraphs):
        replace_docx_paragraph_text(paragraph, replacements[item["id"]])

    if uses_trial:
        complete_study_work_trial(db, current_user.id)
    else:
        increment_usage(db, current_user.id, word_count)
    logger.info("study_work_docx: %s words processed in %.2f seconds", word_count, perf_counter() - started_at)
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return StreamingResponse(
        output,
        media_type=DOCX_MEDIA_TYPE,
        headers={"Content-Disposition": 'attachment; filename="study-work-ready.docx"'},
    )
